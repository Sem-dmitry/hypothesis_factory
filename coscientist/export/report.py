"""
Render a ranked list of :class:`HypothesisAssessment` into business reports.

Four formats are supported, all offline and dependency-light:
- Markdown (plain text)
- self-contained HTML (inline CSS, no external hosts) — printable to PDF
- DOCX (python-docx)
- PDF (reportlab)

Each hypothesis is rendered with the full case-required breakdown: justification,
mechanism of influence, novelty, technical/economic risks, expected value,
target-KPI impact, verification plan, score and citations.
"""

from __future__ import annotations

import html
import os
from typing import Optional, Sequence

from coscientist.hypothesis_assessment import HypothesisAssessment, rank_assessments

DEFAULT_TITLE = "Фабрика гипотез — ранжированный список гипотез"


def _ordered(assessments: Sequence[HypothesisAssessment]) -> list[HypothesisAssessment]:
    return rank_assessments(list(assessments))


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def render_markdown(
    assessments: Sequence[HypothesisAssessment],
    *,
    goal: str = "",
    title: str = DEFAULT_TITLE,
) -> str:
    ranked = _ordered(assessments)
    lines = [f"# {title}", ""]
    if goal:
        lines += [f"**Цель / Goal:** {goal}", ""]
    lines += [f"**Гипотез: {len(ranked)}**", ""]
    for i, a in enumerate(ranked, start=1):
        lines += [
            f"## {i}. {a.hypothesis}",
            "",
            f"- **Оценка (overall score):** {a.overall_score}  "
            f"(новизна {a.novelty_score}, реализуемость {a.feasibility_score}, "
            f"эффект {a.impact_score}, риск {a.risk_level})",
            f"- **Обоснование:** {a.justification}",
            f"- **Механизм влияния:** {a.mechanism_of_influence}",
        ]
        if a.causal_chain:
            lines.append(f"- **Причинно-следственная связь:** {a.causal_chain}")
        if a.world_practice:
            lines.append(f"- **Мировая/промышленная практика:** {a.world_practice}")
        lines.append(f"- **Новизна:** {a.novelty}")
        if a.novelty_vs_input:
            lines.append(f"- **Новизна относительно входных данных:** {a.novelty_vs_input}")
        lines += [
            f"- **Ожидаемая ценность:** {a.expected_value}",
            f"- **Влияние на KPI:** {a.target_kpi_impact}",
        ]
        if a.economic_estimate:
            lines.append(f"- **Экономическая оценка (прикидка):** {a.economic_estimate}")
        if a.kinetics_note:
            lines.append(f"- **Кинетика:** {a.kinetics_note}")
        if a.constraint_adherence:
            lines.append(f"- **Соблюдение ограничений:** {a.constraint_adherence}")
        if a.constraint_violations:
            lines.append(f"- **⚠ Возможные нарушения ограничений:** {'; '.join(a.constraint_violations)}")
        lines += [
            f"- **Технические риски:** {'; '.join(a.technical_risks) or '—'}",
            f"- **Экономические риски:** {'; '.join(a.economic_risks) or '—'}",
        ]
        if a.verification_plan:
            lines.append("- **Дорожная карта проверки:**")
            for step in a.verification_plan:
                lines.append(f"    - {step}")
        if a.citations:
            lines.append("- **Источники:**")
            for c in a.citations:
                lines.append(f"    - {c}")
        lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# HTML (self-contained)
# ---------------------------------------------------------------------------

_HTML_CSS = """
body{font-family:-apple-system,Segoe UI,Roboto,Arial,sans-serif;max-width:900px;
margin:2rem auto;padding:0 1rem;color:#1a1a1a;line-height:1.5}
h1{border-bottom:3px solid #2b6cb0;padding-bottom:.3rem}
.card{border:1px solid #e2e8f0;border-radius:8px;padding:1rem 1.25rem;margin:1rem 0;
box-shadow:0 1px 3px rgba(0,0,0,.06)}
.score{display:inline-block;background:#2b6cb0;color:#fff;border-radius:6px;
padding:.15rem .55rem;font-weight:600;font-size:.9rem}
.k{color:#4a5568;font-weight:600}
ul{margin:.3rem 0 .6rem 1.1rem}
.src{color:#2b6cb0;font-size:.88rem}
.meta{color:#718096;font-size:.85rem}
"""


def _esc(text: str) -> str:
    return html.escape(str(text or ""))


def render_html(
    assessments: Sequence[HypothesisAssessment],
    *,
    goal: str = "",
    title: str = DEFAULT_TITLE,
) -> str:
    ranked = _ordered(assessments)
    parts = [
        "<!doctype html><html lang='ru'><head><meta charset='utf-8'>",
        f"<title>{_esc(title)}</title><style>{_HTML_CSS}</style></head><body>",
        f"<h1>{_esc(title)}</h1>",
    ]
    if goal:
        parts.append(f"<p class='meta'><b>Цель:</b> {_esc(goal)}</p>")
    parts.append(f"<p class='meta'>Гипотез: {len(ranked)}</p>")
    for i, a in enumerate(ranked, start=1):
        parts.append("<div class='card'>")
        parts.append(
            f"<h2>{i}. {_esc(a.hypothesis)} "
            f"<span class='score'>{a.overall_score}</span></h2>"
        )
        parts.append(
            f"<p class='meta'>новизна {a.novelty_score} · реализуемость "
            f"{a.feasibility_score} · эффект {a.impact_score} · риск {a.risk_level}</p>"
        )
        for label, val in [
            ("Обоснование", a.justification),
            ("Механизм влияния", a.mechanism_of_influence),
            ("Новизна", a.novelty),
            ("Ожидаемая ценность", a.expected_value),
            ("Влияние на KPI", a.target_kpi_impact),
        ]:
            if val:
                parts.append(f"<p><span class='k'>{label}:</span> {_esc(val)}</p>")
        if a.technical_risks:
            parts.append("<p><span class='k'>Технические риски:</span></p><ul>")
            parts += [f"<li>{_esc(r)}</li>" for r in a.technical_risks]
            parts.append("</ul>")
        if a.economic_risks:
            parts.append("<p><span class='k'>Экономические риски:</span></p><ul>")
            parts += [f"<li>{_esc(r)}</li>" for r in a.economic_risks]
            parts.append("</ul>")
        if a.verification_plan:
            parts.append("<p><span class='k'>Дорожная карта проверки:</span></p><ul>")
            parts += [f"<li>{_esc(s)}</li>" for s in a.verification_plan]
            parts.append("</ul>")
        if a.citations:
            parts.append("<p><span class='k'>Источники:</span></p><ul class='src'>")
            parts += [f"<li>{_esc(c)}</li>" for c in a.citations]
            parts.append("</ul>")
        parts.append("</div>")
    parts.append("</body></html>")
    return "".join(parts)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def render_docx(
    assessments: Sequence[HypothesisAssessment],
    path: str,
    *,
    goal: str = "",
    title: str = DEFAULT_TITLE,
) -> str:
    from docx import Document

    ranked = _ordered(assessments)
    doc = Document()
    doc.add_heading(title, level=0)
    if goal:
        doc.add_paragraph(f"Цель: {goal}")
    doc.add_paragraph(f"Гипотез: {len(ranked)}")

    for i, a in enumerate(ranked, start=1):
        doc.add_heading(f"{i}. {a.hypothesis}", level=1)
        doc.add_paragraph(
            f"Оценка: {a.overall_score} (новизна {a.novelty_score}, реализуемость "
            f"{a.feasibility_score}, эффект {a.impact_score}, риск {a.risk_level})"
        )
        for label, val in [
            ("Обоснование", a.justification),
            ("Механизм влияния", a.mechanism_of_influence),
            ("Новизна", a.novelty),
            ("Ожидаемая ценность", a.expected_value),
            ("Влияние на KPI", a.target_kpi_impact),
        ]:
            if val:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(val))
        for label, items in [
            ("Технические риски", a.technical_risks),
            ("Экономические риски", a.economic_risks),
            ("Дорожная карта проверки", a.verification_plan),
            ("Источники", a.citations),
        ]:
            if items:
                doc.add_paragraph(label + ":")
                for it in items:
                    doc.add_paragraph(str(it), style="List Bullet")

    _ensure_parent(path)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# PDF (reportlab)
# ---------------------------------------------------------------------------


def render_pdf(
    assessments: Sequence[HypothesisAssessment],
    path: str,
    *,
    goal: str = "",
    title: str = DEFAULT_TITLE,
) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    ranked = _ordered(assessments)
    styles = getSampleStyleSheet()
    _ensure_parent(path)
    doc = SimpleDocTemplate(path, pagesize=A4, title=title)
    flow = [Paragraph(_esc(title), styles["Title"])]
    if goal:
        flow.append(Paragraph(f"<b>Goal:</b> {_esc(goal)}", styles["Normal"]))
    flow.append(Paragraph(f"Hypotheses: {len(ranked)}", styles["Normal"]))
    flow.append(Spacer(1, 12))

    for i, a in enumerate(ranked, start=1):
        flow.append(Paragraph(f"{i}. {_esc(a.hypothesis)}", styles["Heading2"]))
        flow.append(
            Paragraph(
                f"Score: {a.overall_score} (novelty {a.novelty_score}, "
                f"feasibility {a.feasibility_score}, impact {a.impact_score}, "
                f"risk {a.risk_level})",
                styles["Italic"],
            )
        )
        for label, val in [
            ("Justification", a.justification),
            ("Mechanism", a.mechanism_of_influence),
            ("Novelty", a.novelty),
            ("Expected value", a.expected_value),
            ("KPI impact", a.target_kpi_impact),
        ]:
            if val:
                flow.append(Paragraph(f"<b>{label}:</b> {_esc(val)}", styles["Normal"]))
        for label, items in [
            ("Technical risks", a.technical_risks),
            ("Economic risks", a.economic_risks),
            ("Verification plan", a.verification_plan),
            ("Sources", a.citations),
        ]:
            if items:
                flow.append(Paragraph(f"<b>{label}:</b>", styles["Normal"]))
                flow.append(
                    ListFlowable(
                        [ListItem(Paragraph(_esc(it), styles["Normal"])) for it in items],
                        bulletType="bullet",
                    )
                )
        flow.append(Spacer(1, 10))

    doc.build(flow)
    return path


# ---------------------------------------------------------------------------
# Convenience
# ---------------------------------------------------------------------------


def write_report(
    assessments: Sequence[HypothesisAssessment],
    out_dir: str,
    *,
    goal: str = "",
    title: str = DEFAULT_TITLE,
    basename: str = "hypotheses_report",
    formats: Optional[Sequence[str]] = None,
) -> dict[str, str]:
    """Write the report in the requested formats; returns {format: path}."""
    formats = list(formats) if formats else ["md", "html", "docx", "pdf"]
    os.makedirs(out_dir, exist_ok=True)
    written: dict[str, str] = {}
    if "md" in formats:
        p = os.path.join(out_dir, basename + ".md")
        _write_text(p, render_markdown(assessments, goal=goal, title=title))
        written["md"] = p
    if "html" in formats:
        p = os.path.join(out_dir, basename + ".html")
        _write_text(p, render_html(assessments, goal=goal, title=title))
        written["html"] = p
    if "docx" in formats:
        written["docx"] = render_docx(
            assessments, os.path.join(out_dir, basename + ".docx"), goal=goal, title=title
        )
    if "pdf" in formats:
        written["pdf"] = render_pdf(
            assessments, os.path.join(out_dir, basename + ".pdf"), goal=goal, title=title
        )
    return written


def _ensure_parent(path: str) -> None:
    parent = os.path.dirname(os.path.abspath(path))
    os.makedirs(parent, exist_ok=True)


def _write_text(path: str, text: str) -> None:
    _ensure_parent(path)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(text)
