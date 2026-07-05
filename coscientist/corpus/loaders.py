"""
Document loaders for the private corpus.

Each supported file type is turned into a list of :class:`CorpusChunk` records
that carry the text plus enough metadata to build a precise, verifiable citation
(source name + locator + modality). Loaders are pure/offline for PDF/DOCX/XLSX;
image loading delegates text extraction to an injectable VLM callable.
"""

from __future__ import annotations

import hashlib
import os
from dataclasses import dataclass, field
from typing import Callable, Iterable, Iterator, Optional

# Modality -> file extensions it handles.
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}

DEFAULT_CHUNK_SIZE = 1200
DEFAULT_CHUNK_OVERLAP = 150


@dataclass
class CorpusChunk:
    """A single retrievable, citable unit of the corpus."""

    text: str
    source_path: str
    source_name: str
    modality: str  # "pdf" | "docx" | "xlsx" | "image"
    locator: str  # e.g. "p.12", "sheet 'Итог' rows 2-20", "image"
    chunk_id: str = ""
    metadata: dict = field(default_factory=dict)

    def __post_init__(self) -> None:
        if not self.chunk_id:
            self.chunk_id = self._make_id()

    def _make_id(self) -> str:
        digest = hashlib.sha1(
            f"{self.source_path}|{self.locator}|{self.text[:64]}".encode("utf-8")
        ).hexdigest()[:16]
        return f"{self.source_name}:{self.locator}:{digest}"


# ---------------------------------------------------------------------------
# Text chunking
# ---------------------------------------------------------------------------


def chunk_text(
    text: str,
    *,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[str]:
    """Split text into overlapping chunks on paragraph/whitespace boundaries."""
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        if end < n:
            # Prefer to break on a paragraph, then a sentence, then whitespace.
            window = text[start:end]
            for sep in ("\n\n", "\n", ". ", " "):
                idx = window.rfind(sep)
                if idx > chunk_size * 0.5:
                    end = start + idx + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks


# ---------------------------------------------------------------------------
# Per-modality loaders
# ---------------------------------------------------------------------------


def load_pdf(path: str, *, chunk_size: int = DEFAULT_CHUNK_SIZE) -> list[CorpusChunk]:
    """Extract text page-by-page from a PDF using PyMuPDF and chunk it."""
    import fitz  # PyMuPDF

    source_name = os.path.basename(path)
    chunks: list[CorpusChunk] = []
    with fitz.open(path) as doc:
        for page_index in range(doc.page_count):
            page_text = doc[page_index].get_text("text") or ""
            page_no = page_index + 1
            for piece in chunk_text(page_text, chunk_size=chunk_size):
                chunks.append(
                    CorpusChunk(
                        text=piece,
                        source_path=path,
                        source_name=source_name,
                        modality="pdf",
                        locator=f"p.{page_no}",
                        metadata={"page": page_no},
                    )
                )
    return chunks


def load_docx(path: str, *, chunk_size: int = DEFAULT_CHUNK_SIZE) -> list[CorpusChunk]:
    """Extract paragraphs and table cells from a DOCX and chunk them."""
    from docx import Document

    source_name = os.path.basename(path)
    doc = Document(path)
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    full_text = "\n".join(parts)
    return [
        CorpusChunk(
            text=piece,
            source_path=path,
            source_name=source_name,
            modality="docx",
            locator="document",
            metadata={},
        )
        for piece in chunk_text(full_text, chunk_size=chunk_size)
    ]


def load_xlsx(path: str, *, max_rows_per_chunk: int = 40) -> list[CorpusChunk]:
    """Serialize each worksheet to pipe-delimited rows, chunked by row blocks."""
    import openpyxl
    from coscientist.tailings_domain import detect_header_lines, enrich_xlsx_text

    source_name = os.path.basename(path)
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    chunks: list[CorpusChunk] = []
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[tuple[int, str]] = []
            for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                cells = ["" if v is None else str(v) for v in row]
                line = " | ".join(cells).strip(" |")
                if line.strip():
                    rows.append((r_idx, line))
            header_lines = detect_header_lines(rows)
            for block_start in range(0, len(rows), max_rows_per_chunk):
                block = rows[block_start : block_start + max_rows_per_chunk]
                if not block:
                    continue
                first_row, last_row = block[0][0], block[-1][0]
                raw_text = "\n".join(line for _, line in block)
                text, metadata = enrich_xlsx_text(
                    raw_text=raw_text,
                    source_name=source_name,
                    sheet_name=sheet_name,
                    row_start=first_row,
                    row_end=last_row,
                    header_lines=header_lines,
                )
                chunks.append(
                    CorpusChunk(
                        text=text,
                        source_path=path,
                        source_name=source_name,
                        modality="xlsx",
                        locator=f"sheet '{sheet_name}' rows {first_row}-{last_row}",
                        metadata=metadata,
                    )
                )
    finally:
        wb.close()
    return chunks


def load_image(
    path: str,
    *,
    vlm_fn: Optional[Callable[[str], str]] = None,
) -> list[CorpusChunk]:
    """
    Turn an image into a text CorpusChunk via a vision model.

    Parameters
    ----------
    vlm_fn : callable(path) -> str
        Injected image-to-text function. Defaults to
        :func:`coscientist.corpus.vlm.describe_image`. Injecting a stub keeps
        this offline and testable.
    """
    if vlm_fn is None:
        from coscientist.corpus.vlm import describe_image

        vlm_fn = describe_image

    source_name = os.path.basename(path)
    description = (vlm_fn(path) or "").strip()
    if not description:
        return []
    return [
        CorpusChunk(
            text=description,
            source_path=path,
            source_name=source_name,
            modality="image",
            locator="image",
            metadata={"vlm": True},
        )
    ]


# ---------------------------------------------------------------------------
# Dispatch / walking
# ---------------------------------------------------------------------------


def modality_for(path: str) -> Optional[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext in PDF_EXTS:
        return "pdf"
    if ext in DOCX_EXTS:
        return "docx"
    if ext in XLSX_EXTS:
        return "xlsx"
    if ext in IMAGE_EXTS:
        return "image"
    return None


def load_file(
    path: str,
    *,
    include_images: bool = True,
    vlm_fn: Optional[Callable[[str], str]] = None,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
) -> list[CorpusChunk]:
    """Load a single file into CorpusChunks based on its extension."""
    modality = modality_for(path)
    if modality == "pdf":
        return load_pdf(path, chunk_size=chunk_size)
    if modality == "docx":
        return load_docx(path, chunk_size=chunk_size)
    if modality == "xlsx":
        return load_xlsx(path)
    if modality == "image":
        if not include_images:
            return []
        return load_image(path, vlm_fn=vlm_fn)
    return []


def iter_corpus_files(data_dir: str) -> Iterator[str]:
    """Yield supported file paths under ``data_dir`` (recursively, sorted)."""
    for root, _dirs, files in os.walk(data_dir):
        for name in sorted(files):
            full = os.path.join(root, name)
            if modality_for(full) is not None:
                yield full


def load_directory(
    data_dir: str,
    *,
    include_images: bool = True,
    vlm_fn: Optional[Callable[[str], str]] = None,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    files: Optional[Iterable[str]] = None,
) -> list[CorpusChunk]:
    """Load every supported file under ``data_dir`` (or an explicit file list)."""
    paths = list(files) if files is not None else list(iter_corpus_files(data_dir))
    all_chunks: list[CorpusChunk] = []
    for path in paths:
        all_chunks.extend(
            load_file(
                path,
                include_images=include_images,
                vlm_fn=vlm_fn,
                chunk_size=chunk_size,
            )
        )
    return all_chunks
